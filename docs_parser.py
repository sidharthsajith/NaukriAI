import os
import json
from typing import Dict, Union, Optional
from pathlib import Path

# Third-party imports
try:
    import PyPDF2
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:
    print("Required packages not found. Installing...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "PyPDF2", "python-docx"])
    import PyPDF2
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError


class DocumentParser:
    """
    A class to parse documents (PDF and DOCX) and extract text content in a structured format.
    """
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx']
    
    def _is_pdf(self, file_path: Union[str, Path]) -> bool:
        """Check if the file is a PDF."""
        return str(file_path).lower().endswith('.pdf')
    
    def _is_docx(self, file_path: Union[str, Path]) -> bool:
        """Check if the file is a DOCX."""
        return str(file_path).lower().endswith('.docx')
    
    def _extract_text_from_pdf(self, file_path: Union[str, Path]) -> Dict:
        """Extract text from a PDF file."""
        result = {
            'metadata': {},
            'pages': [],
            'total_pages': 0,
            'text': ''
        }
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                result['metadata'] = {
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'creator': pdf_reader.metadata.get('/Creator', ''),
                    'producer': pdf_reader.metadata.get('/Producer', ''),
                    'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                    'modification_date': str(pdf_reader.metadata.get('/ModDate', '')),
                }
                
                result['total_pages'] = len(pdf_reader.pages)
                
                for page_num in range(result['total_pages']):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text() or ''
                    result['pages'].append({
                        'page_number': page_num + 1,
                        'text': page_text.strip()
                    })
                    result['text'] += f"{page_text}\n"
                
                result['text'] = result['text'].strip()
                
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
        
        return result
    
    def _extract_text_from_docx(self, file_path: Union[str, Path]) -> Dict:
        """Extract text from a DOCX file."""
        result = {
            'metadata': {},
            'paragraphs': [],
            'tables': [],
            'text': ''
        }
        
        try:
            doc = Document(file_path)
            
            # Extract core properties
            core_props = doc.core_properties
            result['metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'last_modified_by': core_props.last_modified_by or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'revision': str(core_props.revision) if core_props.revision else '',
                'category': core_props.category or '',
                'keywords': core_props.keywords or '',
                'subject': core_props.subject or ''
            }
            
            # Extract paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    result['paragraphs'].append({
                        'index': i + 1,
                        'text': para.text.strip(),
                        'style': para.style.name if para.style else ''
                    })
                    result['text'] += f"{para.text}\n"
            
            # Extract tables
            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                for i, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                if table_data:
                    result['tables'].append({
                        'table_number': table_num,
                        'data': table_data
                    })
            
            result['text'] = result['text'].strip()
            
        except PackageNotFoundError:
            raise Exception("The file is not a valid DOCX document or is corrupted.")
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
        
        return result
    
    def parse_document(self, file_path: Union[str, Path]) -> Dict:
        """
        Parse a document and return its content in a structured format.
        
        Args:
            file_path: Path to the document file (PDF or DOCX)
            
        Returns:
            dict: Structured content of the document
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            ValueError: If the file format is not supported
            Exception: For other parsing errors
        """
        file_path = Path(file_path)
        
        # Check if file exists
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file extension
        if not (self._is_pdf(file_path) or self._is_docx(file_path)):
            raise ValueError(f"Unsupported file format. Supported formats: {', '.join(self.supported_formats)}")
        
        # Parse based on file type
        if self._is_pdf(file_path):
            return self._extract_text_from_pdf(file_path)
        else:  # DOCX
            return self._extract_text_from_docx(file_path)
    
    def to_json(self, file_path: Union[str, Path], output_file: Optional[Union[str, Path]] = None) -> str:
        """
        Parse a document and save the result as a JSON file.
        
        Args:
            file_path: Path to the input document
            output_file: Path to save the JSON output. If None, uses input filename with .json extension
            
        Returns:
            str: JSON string of the parsed document
        """
        result = self.parse_document(file_path)
        json_str = json.dumps(result, indent=2, ensure_ascii=False)
        
        if output_file is None:
            output_file = Path(file_path).with_suffix('.json')
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(json_str)
        
        return json_str


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description='Parse PDF or DOCX files and extract text in JSON format.')
    parser.add_argument('file_path', help='Path to the PDF or DOCX file')
    parser.add_argument('-o', '--output', help='Output JSON file path (optional)')
    
    args = parser.parse_args()
    
    try:
        doc_parser = DocumentParser()
        json_result = doc_parser.to_json(args.file_path, args.output)
        print("Document parsed successfully!")
        if not args.output:
            print("\nParsed content:")
            print(json_result)
        else:
            print(f"Output saved to: {args.output}")
    except Exception as e:
        print(f"Error: {str(e)}")
        return 1
    
    return 0


if __name__ == "__main__":
    import sys
    sys.exit(main())
